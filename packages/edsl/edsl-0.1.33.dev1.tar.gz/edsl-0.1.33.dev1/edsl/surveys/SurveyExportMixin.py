"""A mixin class for exporting surveys to different formats."""

from typing import Union, Optional


class SurveyExportMixin:
    """A mixin class for exporting surveys to different formats."""

    def css(self):
        from edsl.surveys.SurveyCSS import SurveyCSS

        return SurveyCSS.default_style().generate_css()

    def get_description(self) -> str:
        """Return the description of the survey."""
        from edsl import QuestionFreeText

        question_texts = "\n".join([q.question_text for q in self._questions])
        q = QuestionFreeText(
            question_name="description",
            question_text=f"""A survey was conducted with the following questions: 
                             {question_texts}
                             Please write a description of the survey.
                             """,
        )
        return q.run().select("description").first()

    def docx(self, filename=None) -> Union["Document", None]:
        """Generate a docx document for the survey."""
        from docx import Document

        doc = Document()
        doc.add_heading("EDSL Survey")
        doc.add_paragraph(f"\n")
        for index, question in enumerate(self._questions):
            h = doc.add_paragraph()  # Add question as a paragraph
            h.add_run(f"Question {index + 1} ({question.question_name})").bold = True
            h.add_run(f"; {question.question_type}").italic = True
            p = doc.add_paragraph()
            p.add_run(question.question_text)
            if question.question_type == "linear_scale":
                for key, value in getattr(question, "option_labels", {}).items():
                    doc.add_paragraph(str(key) + ": " + str(value), style="ListBullet")
            else:
                if hasattr(question, "question_options"):
                    for option in getattr(question, "question_options", []):
                        doc.add_paragraph(str(option), style="ListBullet")
        if filename:
            doc.save(filename)
            print("The survey has been saved to", filename)
            return
        return doc

    def to_scenario_list(self) -> "ScenarioList":
        from edsl import ScenarioList, Scenario

        all_keys = set([])
        scenarios = ScenarioList()
        for q in self._questions:
            d = q.to_dict()
            all_keys.update(d.keys())
            scenarios.append(Scenario(d))

        for scenario in scenarios:
            for key in all_keys:
                if key not in scenario:
                    scenario[key] = None

        return scenarios

    def code(self, filename: str = None, survey_var_name: str = "survey") -> list[str]:
        """Create the Python code representation of a survey.

        :param filename: The name of the file to save the code to.
        :param survey_var_name: The name of the survey variable.

        >>> from edsl.surveys import Survey
        >>> survey = Survey.example()
        >>> print(survey.code())
        from edsl.surveys.Survey import Survey
        ...
        ...
        survey = Survey(questions=[q0, q1, q2])
        ...
        """
        import black

        header_lines = ["from edsl.surveys.Survey import Survey"]
        header_lines.append("from edsl import Question")
        lines = ["\n".join(header_lines)]
        for question in self._questions:
            question.question_text = question["question_text"].replace("\n", " ")
            # remove dublicate spaces
            question.question_text = " ".join(question.question_text.split())
            lines.append(f"{question.question_name} = " + repr(question))
        lines.append(
            f"{survey_var_name} = Survey(questions = [{', '.join(self.question_names)}])"
        )
        # return lines
        code_string = "\n".join(lines)
        formatted_code = black.format_str(code_string, mode=black.FileMode())

        if filename:
            print("The code has been saved to", filename)
            print("The survey itself is saved to 'survey' object")
            with open(filename, "w") as file:
                file.write(formatted_code)
            return

        return formatted_code

    def html(
        self,
        scenario: Optional[dict] = None,
        filename: Optional[str] = None,
        return_link=False,
        css: Optional[str] = None,
        cta: Optional[str] = "Open HTML file",
        include_question_name=False,
    ):
        from IPython.display import display, HTML
        import tempfile
        import os
        from edsl.utilities.utilities import is_notebook

        if scenario is None:
            scenario = {}

        if css is None:
            css = self.css()

        if filename is None:
            current_directory = os.getcwd()
            filename = tempfile.NamedTemporaryFile(
                "w", delete=False, suffix=".html", dir=current_directory
            ).name

        html_header = f"""<html>
        <head><title></title>
        <style>
        { css }
        </style>
        </head>
        <body>
        <div class="survey_container">
        """

        html_footer = """
        </div>
        </body>
        </html>"""

        output = html_header

        with open(filename, "w") as f:
            f.write(html_header)
            for question in self._questions:
                f.write(
                    question.html(
                        scenario=scenario, include_question_name=include_question_name
                    )
                )
                output += question.html(
                    scenario=scenario, include_question_name=include_question_name
                )
            f.write(html_footer)
            output += html_footer

        if is_notebook():
            html_url = f"/files/{filename}"
            html_link = f'<a href="{html_url}" target="_blank">{cta}</a>'
            display(HTML(html_link))

            import html

            escaped_output = html.escape(output)
            iframe = f""""
            <iframe srcdoc="{ escaped_output }" style="width: 800px; height: 600px;"></iframe>
            """
            display(HTML(iframe))

        else:
            print(f"Survey saved to {filename}")
            import webbrowser
            import os

            webbrowser.open(f"file://{os.path.abspath(filename)}")
            # webbrowser.open(filename)

        if return_link:
            return filename


if __name__ == "__main__":
    import doctest

    doctest.testmod(optionflags=doctest.ELLIPSIS)
